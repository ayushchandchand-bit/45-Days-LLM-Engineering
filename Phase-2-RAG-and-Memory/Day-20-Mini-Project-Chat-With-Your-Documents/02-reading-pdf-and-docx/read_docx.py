"""
Day 20 - Step 2a: Read text out of a Word (.docx) file with python-docx.

Install once:
    pip install python-docx     # note: pip name is python-docx, import name is docx

Run it:
    python read_docx.py

The script writes a sample resume.docx next to itself, then reads it back.
"""

from pathlib import Path

# The pip package is "python-docx" but you import it as "docx".
from docx import Document

# The sample file lives next to this script.
SAMPLE = Path(__file__).resolve().parent / "sample_resume.docx"


def make_sample_docx() -> None:
    """Create a small .docx so this demo needs no setup. (Ignore the details;
    the part you reuse is read_docx() below.)"""
    doc = Document()
    doc.add_heading("Priya Sharma", level=0)
    doc.add_paragraph("B.Tech (CSE), 2026 - Pune")
    doc.add_paragraph("Skills: Python, SQL, Streamlit, basic ML")
    doc.add_paragraph("Built a semantic search engine over lecture notes.")
    doc.save(SAMPLE)


def read_docx(path) -> str:
    """The 3 lines that matter: open the doc, read each paragraph's text, join."""
    # A Word document is a sequence of paragraphs. Document(path) opens it.
    doc = Document(path)
    # Each paragraph has a .text attribute. Join them with newlines into one string.
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def main() -> None:
    make_sample_docx()

    print("=" * 72)
    print("Reading a .docx file with python-docx")
    print("=" * 72)
    print(f"File: {SAMPLE.name}\n")

    text = read_docx(SAMPLE)
    print("Extracted text:")
    print(text)
    print(f"\n({len(text)} characters extracted)")


if __name__ == "__main__":
    main()
